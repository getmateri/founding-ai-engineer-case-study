"""
Extraction Logic

Handles parsing source files (Excel, Markdown, ZIP archives) and extracting term sheet fields using LLM.
"""

import json
import os
import logging
import zipfile
import tempfile
from pathlib import Path
from typing import Optional
from openai import OpenAI
from openpyxl import load_workbook

from .schema import (
    TermSheetData, ExtractedField, SourceReference, ConflictingValue,
    PartiesSection, DealEconomicsSection, LiquidationTermsSection,
    GovernanceSection, FounderTermsSection, TransactionTermsSection, SignaturesSection
)

logger = logging.getLogger("extraction")

# Optional imports for document parsing
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed, .docx files won't be parsed")

try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False
    logger.warning("pypdf not installed, .pdf files won't be parsed")


# =============================================================================
# FILE PARSING
# =============================================================================

def parse_excel(file_path: str) -> str:
    """Parse Excel file into text representation"""
    logger.info(f"Parsing Excel file: {file_path}")
    try:
        wb = load_workbook(file_path, data_only=True)
        logger.debug(f"Found {len(wb.sheetnames)} sheets: {wb.sheetnames}")
    except Exception as e:
        logger.error(f"Failed to load Excel file: {e}")
        raise
    
    content_parts = []
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        content_parts.append(f"\n=== Sheet: {sheet_name} ===\n")
        
        row_count = 0
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            # Skip empty rows
            if all(cell is None for cell in row):
                continue
            
            row_content = []
            for col_idx, cell in enumerate(row, start=1):
                if cell is not None:
                    col_letter = _col_num_to_letter(col_idx)
                    row_content.append(f"{col_letter}{row_idx}: {cell}")
            
            if row_content:
                content_parts.append(" | ".join(row_content))
                row_count += 1
        
        logger.debug(f"  Sheet '{sheet_name}': {row_count} non-empty rows")
    
    result = "\n".join(content_parts)
    logger.info(f"Excel parsed: {len(result)} characters")
    return result


def _col_num_to_letter(col_num: int) -> str:
    """Convert column number to Excel letter (1=A, 2=B, etc.)"""
    result = ""
    while col_num > 0:
        col_num, remainder = divmod(col_num - 1, 26)
        result = chr(65 + remainder) + result
    return result


def parse_markdown(file_path: str) -> str:
    """Read markdown file"""
    with open(file_path, 'r') as f:
        return f.read()


def parse_text_file(file_path: str) -> str:
    """Read plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()


def parse_docx(file_path: str) -> str:
    """Parse .docx file into text"""
    if not HAS_DOCX:
        logger.warning(f"Cannot parse {file_path}: python-docx not installed")
        return f"[Could not parse .docx file: {file_path}]"
    
    logger.info(f"Parsing DOCX file: {file_path}")
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    
    result = "\n\n".join(paragraphs)
    logger.info(f"DOCX parsed: {len(result)} characters")
    return result


def parse_pdf(file_path: str) -> str:
    """Parse .pdf file into text"""
    if not HAS_PYPDF:
        logger.warning(f"Cannot parse {file_path}: pypdf not installed")
        return f"[Could not parse .pdf file: {file_path}]"
    
    logger.info(f"Parsing PDF file: {file_path}")
    reader = pypdf.PdfReader(file_path)
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            pages.append(f"--- Page {i+1} ---\n{text}")
    
    result = "\n\n".join(pages)
    logger.info(f"PDF parsed: {len(result)} characters from {len(reader.pages)} pages")
    return result


def parse_doc_legacy(file_path: str) -> str:
    """
    Parse legacy .doc file. These are binary files that are harder to parse.
    We'll try to extract readable text using a simple approach.
    """
    logger.info(f"Parsing legacy .doc file: {file_path}")
    try:
        # Try to read as binary and extract text
        with open(file_path, 'rb') as f:
            content = f.read()
        
        # Legacy .doc files have text interspersed with binary data
        # Try to extract readable ASCII/UTF-8 portions
        text_parts = []
        current_text = []
        
        for byte in content:
            # Check if it's a printable ASCII character or common whitespace
            if 32 <= byte <= 126 or byte in (9, 10, 13):
                current_text.append(chr(byte))
            else:
                if len(current_text) > 20:  # Only keep substantial text chunks
                    text_parts.append(''.join(current_text))
                current_text = []
        
        if len(current_text) > 20:
            text_parts.append(''.join(current_text))
        
        result = '\n'.join(text_parts)
        logger.info(f"Legacy .doc parsed: {len(result)} characters extracted")
        return result
        
    except Exception as e:
        logger.error(f"Failed to parse legacy .doc: {e}")
        return f"[Could not parse legacy .doc file: {file_path}]"


def parse_file(file_path: str) -> Optional[str]:
    """Parse a file based on its extension"""
    path = Path(file_path)
    suffix = path.suffix.lower()
    
    parsers = {
        '.xlsx': parse_excel,
        '.xls': parse_excel,
        '.md': parse_markdown,
        '.txt': parse_text_file,
        '.docx': parse_docx,
        '.pdf': parse_pdf,
        '.doc': parse_doc_legacy,
    }
    
    parser = parsers.get(suffix)
    if parser:
        try:
            return parser(file_path)
        except Exception as e:
            logger.error(f"Failed to parse {file_path}: {e}")
            return None
    else:
        logger.warning(f"Unknown file type: {suffix} for {file_path}")
        return None


def extract_zip(zip_path: str, extract_to: str) -> list[str]:
    """Extract zip file and return list of extracted file paths"""
    logger.info(f"Extracting ZIP: {zip_path}")
    extracted_files = []
    
    with zipfile.ZipFile(zip_path, 'r') as zf:
        # List contents
        for info in zf.infolist():
            if not info.is_dir():
                logger.debug(f"  ZIP contains: {info.filename}")
        
        # Extract all
        zf.extractall(extract_to)
        
        # Get list of extracted files
        for info in zf.infolist():
            if not info.is_dir():
                extracted_path = os.path.join(extract_to, info.filename)
                if os.path.exists(extracted_path):
                    extracted_files.append(extracted_path)
    
    logger.info(f"Extracted {len(extracted_files)} files from ZIP")
    return extracted_files


def load_source_files(data_dir: str = "data") -> dict[str, str]:
    """Load all source files from data directory, including ZIP archives"""
    logger.info(f"Loading source files from: {data_dir}")
    sources = {}
    data_path = Path(data_dir)
    
    if not data_path.exists():
        logger.error(f"Data directory not found: {data_path.absolute()}")
        return sources
    
    logger.debug(f"Data directory contents: {list(data_path.iterdir())}")
    
    # Load deal model (Excel)
    excel_path = data_path / "Model.xlsx"
    if excel_path.exists():
        logger.info(f"Loading deal model: {excel_path}")
        sources["deal_model"] = parse_excel(str(excel_path))
    else:
        logger.warning(f"Deal model not found: {excel_path}")
    
    # Load firm policy (Markdown)
    policy_path = data_path / "firm_policy.md"
    if policy_path.exists():
        logger.info(f"Loading firm policy: {policy_path}")
        sources["firm_policy"] = parse_markdown(str(policy_path))
    else:
        logger.warning(f"Firm policy not found: {policy_path}")
    
    # Process ZIP files (e.g., Termsheets.zip)
    for zip_file in data_path.glob("*.zip"):
        logger.info(f"Found ZIP archive: {zip_file}")
        
        # Create temp directory to extract to
        with tempfile.TemporaryDirectory() as temp_dir:
            extracted_files = extract_zip(str(zip_file), temp_dir)
            
            # Parse each extracted file
            for file_path in extracted_files:
                file_name = Path(file_path).name
                
                # Skip hidden files and system files
                if file_name.startswith('.') or file_name.startswith('__'):
                    continue
                
                content = parse_file(file_path)
                if content:
                    # Use a descriptive key based on the zip and file name
                    source_key = f"{zip_file.stem}/{file_name}"
                    sources[source_key] = content
                    logger.info(f"Loaded from ZIP: {source_key} ({len(content)} chars)")
    
    # Also check for a Termsheets directory (if already extracted)
    termsheets_dir = data_path / "Termsheets"
    if termsheets_dir.exists() and termsheets_dir.is_dir():
        logger.info(f"Found Termsheets directory: {termsheets_dir}")
        for file_path in termsheets_dir.iterdir():
            if file_path.is_file() and not file_path.name.startswith('.'):
                content = parse_file(str(file_path))
                if content:
                    source_key = f"Termsheets/{file_path.name}"
                    sources[source_key] = content
                    logger.info(f"Loaded from directory: {source_key} ({len(content)} chars)")
    
    logger.info(f"Loaded {len(sources)} source files: {list(sources.keys())}")
    return sources


# =============================================================================
# LLM EXTRACTION
# =============================================================================

EXTRACTION_SYSTEM_PROMPT = """You are an expert at extracting structured data from financial documents for venture capital term sheets.

Your job is to extract specific fields from the provided source documents. For each field:
1. Find the value in the documents (deal model takes precedence for deal-specific terms, firm policy for standard terms)
2. Note exactly where you found it (file + cell reference or section)
3. Rate your confidence using BINARY scoring:
   - 1.0 = You are 100% certain (value explicitly stated, exact match, no ambiguity)
   - 0.0 = Anything less than 100% certain (inferred, assumed, derived, or any doubt)
4. If multiple sources conflict, note all values

IMPORTANT: Only use confidence 1.0 when the value is EXPLICITLY and UNAMBIGUOUSLY stated. If you had to infer, calculate, or make any assumption, use 0.0 so a human can review.

Be precise with numbers and percentages. Use the exact format requested."""


def extract_section(
    client: OpenAI,
    section_name: str,
    sources: dict[str, str],
    model: str = "gpt-4o"
) -> tuple[dict, dict]:
    """Extract fields for a specific section using LLM
    
    Returns:
        tuple: (extracted_data, token_usage)
    """
    
    section_prompts = {
        "parties": _get_parties_prompt(),
        "deal_economics": _get_deal_economics_prompt(),
        "liquidation_terms": _get_liquidation_terms_prompt(),
        "governance": _get_governance_prompt(),
        "founder_terms": _get_founder_terms_prompt(),
        "transaction_terms": _get_transaction_terms_prompt(),
        "signatures": _get_signatures_prompt(),
    }
    
    prompt = section_prompts.get(section_name)
    if not prompt:
        raise ValueError(f"Unknown section: {section_name}")
    
    # Build the full prompt with all sources
    source_sections = []
    
    # Primary sources - these contain the actual data
    if 'deal_model' in sources:
        source_sections.append(f"=== DEAL MODEL (Model.xlsx) ===\n{sources['deal_model']}")
    
    if 'firm_policy' in sources:
        source_sections.append(f"=== FIRM POLICY (firm_policy.md) ===\n{sources['firm_policy']}")
    
    # Include reference term sheets for formatting guidance
    termsheet_sources = [k for k in sources.keys() if 'Termsheets/' in k or 'termsheet' in k.lower()]
    if termsheet_sources:
        source_sections.append("=== REFERENCE TERM SHEETS (format examples) ===")
        for ts_key in termsheet_sources[:3]:  # Include up to 3
            content = sources[ts_key]
            # Truncate very long docs
            if len(content) > 5000:
                content = content[:5000] + "\n... [truncated]"
            source_sections.append(f"\n--- {ts_key} ---\n{content}")
    
    full_prompt = f"""{prompt}

SOURCE DOCUMENTS:

{chr(10).join(source_sections)}

IMPORTANT:
- Extract values from the DEAL MODEL for deal-specific terms
- Use FIRM POLICY for standard terms and defaults  
- The reference term sheets are for FORMAT guidance only - do NOT extract values from them
- If a value is not found, set value to null and confidence to 0

Respond with valid JSON only, no markdown formatting."""

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": EXTRACTION_SYSTEM_PROMPT},
            {"role": "user", "content": full_prompt}
        ],
        response_format={"type": "json_object"},
        temperature=0.1,
    )
    
    # Track token usage
    token_usage = {
        "section": section_name,
        "input_tokens": response.usage.prompt_tokens if response.usage else 0,
        "output_tokens": response.usage.completion_tokens if response.usage else 0,
    }
    
    return json.loads(response.choices[0].message.content), token_usage


def _get_parties_prompt() -> str:
    return """Extract the PARTIES section fields:

{
  "company_name": {
    "value": "Company name",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "company_jurisdiction": {
    "value": "Delaware/etc",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "founders": {
    "value": "Founder 1, Founder 2",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "lead_investor": {
    "value": "Investor name",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  }
}

If a field is not found, use null for value and set confidence to 0."""


def _get_deal_economics_prompt() -> str:
    return """Extract the DEAL ECONOMICS section fields:

{
  "round_type": {
    "value": "Seed/Series A/etc",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "investment_amount": {
    "value": "$X,XXX,XXX format",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "pre_money_valuation": {
    "value": "$XX,XXX,XXX format",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "Check firm_policy section 2.1 for valuation ranges by round type"
  },
  "security_type": {
    "value": "Series A Preferred Stock/etc",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "price_per_share": {
    "value": "$X.XX or null if not specified",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "target_ownership_pct": {
    "value": "XX%",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 2.2 says target 15-20%"
  },
  "option_pool_pct": {
    "value": "XX%",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 2.3 requires at least 15%"
  },
  "option_pool_timing": {
    "value": "pre-money or post-money",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  }
}

If a field is not found in deal_model, check if firm_policy has a default. Note conflicts if values differ."""


def _get_liquidation_terms_prompt() -> str:
    return """Extract the LIQUIDATION TERMS section fields:

{
  "liquidation_preference_multiple": {
    "value": "1.0 (number only)",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 2.4: must be 1x, never >1x"
  },
  "participation_type": {
    "value": "non-participating/participating/capped participating",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 2.4: always non-participating"
  },
  "dividend_type": {
    "value": "non-cumulative/cumulative/none",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "dividend_rate_pct": {
    "value": "X% or null",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 2.5: 6% for deals under $5M"
  },
  "anti_dilution_type": {
    "value": "broad-based weighted average/narrow-based weighted average/full ratchet/none",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 2.6: must be broad-based weighted average"
  }
}

Use firm_policy defaults where deal_model doesn't specify. Flag any conflicts with policy."""


def _get_governance_prompt() -> str:
    return """Extract the GOVERNANCE section fields:

{
  "board_seats_total": {
    "value": "3 (number only)",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "board_seats_investor": {
    "value": "1 (number only)",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 3.1"
  },
  "board_seats_founder": {
    "value": "2 (number only)",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "board_seats_independent": {
    "value": "0 (number only)",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "board_observer_rights": {
    "value": "true/false",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 3.1: always include observer rights"
  },
  "investor_consent_for_quorum": {
    "value": "true/false",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "drag_along_threshold_pct": {
    "value": "50%",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 3.4"
  },
  "pro_rata_rights": {
    "value": "true/false",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 3.5"
  }
}"""


def _get_founder_terms_prompt() -> str:
    return """Extract the FOUNDER TERMS section fields:

{
  "vesting_period_months": {
    "value": "48 (number only)",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 4.1: 4 years = 48 months"
  },
  "vesting_cliff_months": {
    "value": "12 (number only)",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 4.1: 1 year cliff"
  },
  "vesting_frequency": {
    "value": "monthly/quarterly/annually",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 4.1: monthly"
  },
  "acceleration_type": {
    "value": "none/single-trigger/double-trigger",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 4.1: single-trigger not permitted"
  },
  "non_compete_months": {
    "value": "12 (number only)",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 4.2: 12 months"
  },
  "non_solicit_months": {
    "value": "24 (number only)",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 4.2: 24 months"
  }
}"""


def _get_transaction_terms_prompt() -> str:
    return """Extract the TRANSACTION TERMS section fields:

{
  "exclusivity_days": {
    "value": "45 (number only)",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 5.1: 45 days"
  },
  "legal_fee_cap": {
    "value": "$25,000 format",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "firm_policy section 5.2: $25,000 standard"
  },
  "expected_closing_days": {
    "value": "30 (number only)",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "governing_law": {
    "value": "Delaware/etc",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  }
}"""


def _get_signatures_prompt() -> str:
    return """Extract the SIGNATURES section fields:

Note: Most signature fields will be blank (to be filled at signing). Focus on extracting any mentioned dates or names.

{
  "effective_date": {
    "value": "Date or null if not specified",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "company_signatory_name": {
    "value": "Name of company signatory or null",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "company_signatory_title": {
    "value": "Title (e.g., CEO, CFO) or null",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "investor_signatory_name": {
    "value": "Name of investor signatory or null",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "investor_signatory_title": {
    "value": "Title (e.g., Managing Partner) or null",
    "source": {"file": "...", "location": "..."},
    "confidence": 0.0-1.0,
    "conflicts": [],
    "reasoning": "..."
  },
  "binding_status": {
    "value": "non-binding (term sheets are typically non-binding except exclusivity/confidentiality)",
    "source": {"file": "...", "location": "..."},
    "confidence": 1.0,
    "conflicts": [],
    "reasoning": "Standard term sheet practice"
  }
}"""


# =============================================================================
# MAIN EXTRACTION FUNCTION
# =============================================================================

def extract_all_sections(
    client: OpenAI,
    sources: dict[str, str],
    model: str = "gpt-4o"
) -> tuple[TermSheetData, list[dict]]:
    """Extract all sections and build complete TermSheetData
    
    Returns:
        tuple: (term_sheet_data, llm_calls_list)
    """
    
    term_sheet = TermSheetData()
    llm_calls = []
    
    sections = [
        ("parties", PartiesSection),
        ("deal_economics", DealEconomicsSection),
        ("liquidation_terms", LiquidationTermsSection),
        ("governance", GovernanceSection),
        ("founder_terms", FounderTermsSection),
        ("transaction_terms", TransactionTermsSection),
        ("signatures", SignaturesSection),
    ]
    
    import time
    
    for i, (section_name, section_class) in enumerate(sections):
        # Add delay between requests to avoid rate limiting (except first)
        if i > 0:
            logger.info("Waiting 5s to avoid rate limits...")
            time.sleep(5)
        
        logger.info(f"Extracting section {i+1}/{len(sections)}: {section_name}...")
        raw_data, token_usage = extract_section(client, section_name, sources, model)
        llm_calls.append(token_usage)
        section_data = _parse_extraction_response(raw_data, section_class)
        setattr(term_sheet, section_name, section_data)
    
    return term_sheet, llm_calls


def _parse_extraction_response(raw_data: dict, section_class: type) -> object:
    """Parse LLM response into section model"""
    section = section_class()
    
    for field_name, field_data in raw_data.items():
        if hasattr(section, field_name) and isinstance(field_data, dict):
            # Safely parse source - LLM may return null values
            source = None
            source_data = field_data.get("source")
            if source_data and isinstance(source_data, dict):
                file_val = source_data.get("file")
                location_val = source_data.get("location")
                # Only create SourceReference if both fields are valid strings
                if file_val and location_val:
                    source = SourceReference(file=str(file_val), location=str(location_val))
            
            # Safely parse conflicts
            conflicts = []
            for c in field_data.get("conflicts", []) or []:
                if isinstance(c, dict) and c.get("source") and c.get("value"):
                    try:
                        conflicts.append(ConflictingValue(
                            source=str(c["source"]),
                            value=str(c["value"]),
                            confidence=float(c.get("confidence", 0.5))
                        ))
                    except (ValueError, TypeError):
                        pass  # Skip malformed conflict entries
            
            extracted = ExtractedField(
                value=field_data.get("value"),
                source=source,
                confidence=float(field_data.get("confidence", 0.5)) if field_data.get("confidence") is not None else 0.5,
                conflicts=conflicts,
                found=field_data.get("value") is not None,
                derived_from_policy="firm_policy" in str(source_data or {}),
                reasoning=field_data.get("reasoning"),
            )
            setattr(section, field_name, extracted)
    
    return section

